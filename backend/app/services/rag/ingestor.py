"""
Ingestion des documents : extraction du texte + découpage en chunks.

Le principe pédagogique : on ne met JAMAIS un document entier dans le prompt du LLM
(il est trop long et on perd la précision). On le découpe en "chunks" chevauchants,
chacun sera indexé par un embedding. À la question, on ne retrouve que les chunks
les plus proches de la question.

Formats supportés : .pdf, .docx, .xlsx, .csv, .txt, .md
"""

import csv
import io
import os
from typing import List

from app.config import settings


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: str) -> str:
    import docx

    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _read_xlsx(path: str) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"Feuille: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_csv(path: str) -> str:
    rows = []
    with open(path, "r", encoding="utf-8-sig", errors="replace") as f:
        for row in csv.reader(f):
            rows.append(" | ".join(cells.strip() for cells in row))
    return "\n".join(rows)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def parse_document(path: str) -> str:
    """Extrait tout le texte d'un fichier selon son extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".xlsx":
        return _read_xlsx(path)
    if ext == ".csv":
        return _read_csv(path)
    if ext in (".txt", ".md"):
        return _read_txt(path)
    raise ValueError(f"Format non supporté : {ext}")


def chunk_text(text: str, size: int = 800, overlap: int = 100) -> List[str]:
    """
    Découpe un texte en chunks de `size` caractères avec `overlap` de chevauchement.
    Le chevauchement évite de couper une idée en deux à la frontière de deux chunks.
    """
    text = text.strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            last_space = text.rfind(" ", start, end)
            if last_space > start:
                end = last_space
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def detect_doc_type(filename: str) -> str:
    return os.path.splitext(filename)[1].lstrip(".").lower()


def file_is_supported(filename: str) -> bool:
    return detect_doc_type(filename) in {"pdf", "docx", "xlsx", "csv", "txt", "md"}


def ingest_bytes(filename: str, content: bytes) -> List[str]:
    """Lit un fichier en mémoire (upload) et retourne ses chunks."""
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".csv", ".txt", ".md"):
        text = content.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        import docx

        doc = docx.Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif ext == ".xlsx":
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"Feuille: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    else:
        raise ValueError(f"Format non supporté : {ext}")

    chunks = chunk_text(text, size=settings.AI_TOP_K * 200 + 400, overlap=100)
    return [c for c in chunks if c]
